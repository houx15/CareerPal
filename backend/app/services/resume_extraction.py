from pathlib import Path

import fitz
from docx import Document


PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ResumeExtractionError(Exception):
    pass


def extract_resume_text(path: Path, content_type: str) -> str:
    try:
        if content_type == PDF_CONTENT_TYPE:
            text = _extract_pdf_text(path)
        elif content_type == DOCX_CONTENT_TYPE:
            text = _extract_docx_text(path)
        else:
            raise ResumeExtractionError("Unsupported resume file type")
    except ResumeExtractionError:
        raise
    except Exception as exc:
        raise ResumeExtractionError("Resume file does not contain readable text") from exc

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip()).strip()
    if not normalized:
        raise ResumeExtractionError("Resume file does not contain readable text")
    return normalized


def _extract_pdf_text(path: Path) -> str:
    with fitz.open(path) as document:
        return "\n".join(page.get_text() for page in document)


def _extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)
