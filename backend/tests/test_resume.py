import asyncio
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import fitz
import pytest
from docx import Document
from fastapi import HTTPException

from app.api import resume as resume_api
from app.api.resume import store_validated_resume
from app.core.config import get_settings
from app.models.conversation import Conversation
from app.models.resume import ResumeFile
from app.models.user import Experience, Profile, Project, Skill
from app.services.llm import LLMMessage
from app.services.resume_extraction import extract_resume_text


class StubLLMClient:
    def __init__(self, chunks):
        self.chunks = chunks
        self.messages: list[LLMMessage] | None = None

    async def stream_chat(self, messages):
        self.messages = list(messages)
        for chunk in self.chunks:
            yield chunk


@pytest.fixture(autouse=True)
def clear_settings_cache_after_test():
    yield
    get_settings.cache_clear()


def auth_headers(client):
    response = client.post(
        "/api/auth/register",
        json={"email": "alex.resume@example.com", "username": "alexresume", "password": "secret123"},
    )
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def make_pdf_bytes(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    return document.tobytes()


def make_docx_bytes_with_text(text: str, table_text: str | None = None) -> bytes:
    document = Document()
    document.add_paragraph(text)
    if table_text:
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = table_text
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def configure_resume_storage(monkeypatch, storage_dir, max_upload_bytes=None):
    monkeypatch.setenv("CAREERPAL_RESUME_STORAGE_DIR", str(storage_dir))
    if max_upload_bytes is not None:
        monkeypatch.setenv("CAREERPAL_RESUME_MAX_UPLOAD_BYTES", str(max_upload_bytes))
    get_settings.cache_clear()


def test_extracts_text_from_pdf_fixture(tmp_path):
    path = tmp_path / "resume.pdf"
    path.write_bytes(make_pdf_bytes("Alex Chen Backend Engineer"))

    text = extract_resume_text(path, "application/pdf")

    assert "Alex Chen Backend Engineer" in text


def test_extracts_text_from_docx_fixture(tmp_path):
    path = tmp_path / "resume.docx"
    path.write_bytes(make_docx_bytes_with_text("Built distributed systems", "Python Django PostgreSQL"))

    text = extract_resume_text(
        path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "Built distributed systems" in text
    assert "Python Django PostgreSQL" in text


def test_authenticated_user_can_upload_pdf_resume(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)
    content = make_pdf_bytes("Alex Chen sample resume")

    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", content, "application/pdf")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["id"]
    assert body["original_filename"] == "resume.pdf"
    assert body["content_type"] == "application/pdf"
    assert body["size_bytes"] == len(content)
    assert body["status"] == "parsed"
    assert body["parse_error"] is None
    assert body["parsed_at"]
    assert body["created_at"]
    assert "storage_path" not in body
    resume_file = db_session.get(ResumeFile, body["id"])
    assert resume_file is not None
    assert resume_file.user_id
    stored_path = Path(resume_file.storage_path)
    assert stored_path.exists()
    assert stored_path.read_bytes() == content
    assert stored_path.is_relative_to(tmp_path / "resumes")


def test_resume_upload_requires_authentication(client):
    response = client.post(
        "/api/resume/upload",
        files={"file": ("resume.pdf", b"%PDF-1.4", "application/pdf")},
    )

    assert response.status_code == 401


def test_resume_upload_rejects_unsupported_file_type(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)

    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.txt", b"plain text", "text/plain")},
    )

    assert response.status_code == 415
    assert response.json()["detail"] == "Unsupported resume file type. Upload a PDF or DOCX file."
    assert not (tmp_path / "resumes").exists()
    assert db_session.query(ResumeFile).count() == 0


def test_resume_upload_rejects_oversized_file(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes", max_upload_bytes=10)
    headers = auth_headers(client)

    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", b"%PDF-1.4 too large", "application/pdf")},
    )

    assert response.status_code == 413
    assert response.json()["detail"] == "Resume file is too large"
    assert not (tmp_path / "resumes").exists()
    assert db_session.query(ResumeFile).count() == 0


def test_resume_upload_rejects_empty_file(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)

    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", b"", "application/pdf")},
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "Resume file is empty"
    assert not (tmp_path / "resumes").exists()
    assert db_session.query(ResumeFile).count() == 0


def test_authenticated_user_can_upload_docx_resume(client, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.docx", make_docx_bytes_with_text("Built APIs"), content_type)},
    )

    assert response.status_code == 201
    assert response.json()["content_type"] == content_type
    assert response.json()["original_filename"] == "resume.docx"
    assert response.json()["status"] == "parsed"


def test_upload_extracts_pdf_text_and_exposes_parse_result(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)

    upload = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", make_pdf_bytes("Alex Chen Backend Engineer"), "application/pdf")},
    )

    assert upload.status_code == 201
    resume_id = upload.json()["id"]
    status_response = client.get(f"/api/resume/parse-status/{resume_id}", headers=headers)
    parsed_response = client.get(f"/api/resume/parsed/{resume_id}", headers=headers)

    assert status_response.status_code == 200
    assert status_response.json()["status"] == "parsed"
    assert status_response.json()["parse_error"] is None
    assert parsed_response.status_code == 200
    assert "Alex Chen Backend Engineer" in parsed_response.json()["parsed_text"]

    resume_file = db_session.get(ResumeFile, resume_id)
    assert resume_file.status == "parsed"
    assert "Alex Chen Backend Engineer" in resume_file.parsed_text
    assert resume_file.parsed_at is not None


def test_upload_records_actionable_error_for_unreadable_resume(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)

    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", b"%PDF-1.4\n%%EOF", "application/pdf")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "parse_failed"
    assert body["parse_error"] == "Resume file does not contain readable text"
    resume_id = body["id"]
    parsed_response = client.get(f"/api/resume/parsed/{resume_id}", headers=headers)
    assert parsed_response.status_code == 422
    assert parsed_response.json()["detail"] == "Resume text has not been extracted"


def test_structures_parsed_resume_into_profile_and_follow_up_conversation(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)
    llm_client = StubLLMClient(
        [
            "{",
            '"profile":{"name":"Alex Chen","headline":"Backend engineer","target_direction":"AI platform roles"},',
            '"experience":[{"company":"Acme","role":"Backend Intern","time":"Summer 2025",',
            '"description":"Built resume parsing APIs","achievements":["Cut parsing latency by 30%"]}],',
            '"skills":[{"name":"Python","category":"Backend","proficiency":"advanced"}],',
            '"follow_up_questions":["Which AI platform roles are you targeting?","What was the user impact at Acme?"]',
            "}",
        ]
    )
    monkeypatch.setattr(resume_api, "build_llm_client", lambda settings: llm_client)

    upload = client.post(
        "/api/resume/upload",
        headers=headers,
        files={
            "file": (
                "resume.pdf",
                make_pdf_bytes("Alex Chen Backend Intern Acme Python AI platform roles"),
                "application/pdf",
            )
        },
    )

    response = client.post(f"/api/resume/structure/{upload.json()['id']}", headers=headers)

    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "structured"
    assert body["structure_error"] is None
    assert body["profile"]["name"] == "Alex Chen"
    assert body["profile"]["experience"][0]["company"] == "Acme"
    assert body["follow_up_questions"] == [
        "Which AI platform roles are you targeting?",
        "What was the user impact at Acme?",
    ]

    profile = db_session.query(Profile).one()
    assert profile.name == "Alex Chen"
    assert profile.headline == "Backend engineer"
    assert db_session.query(Experience).one().achievements == ["Cut parsing latency by 30%"]
    assert db_session.query(Skill).one().name == "Python"

    conversation = db_session.query(Conversation).one()
    assert conversation.context_type == "career"
    assert conversation.focus_node == "resume"
    assert conversation.id == body["conversation_id"]
    assert "I imported your resume" in conversation.messages[0]["content"]
    assert "Which AI platform roles are you targeting?" in conversation.messages[0]["content"]
    assert "resume text JSON" in llm_client.messages[1].content


def test_resume_structure_rejects_unparsed_resume(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)
    upload = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", b"%PDF-1.4\n%%EOF", "application/pdf")},
    )

    response = client.post(f"/api/resume/structure/{upload.json()['id']}", headers=headers)

    assert response.status_code == 422
    assert response.json()["detail"] == "Resume text has not been extracted"
    assert db_session.query(Experience).count() == 0
    assert db_session.query(Skill).count() == 0


def test_resume_structure_records_actionable_error_for_malformed_llm_json(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)
    monkeypatch.setattr(resume_api, "build_llm_client", lambda settings: StubLLMClient(["not json"]))
    upload = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", make_pdf_bytes("Alex Chen Backend Engineer"), "application/pdf")},
    )

    response = client.post(f"/api/resume/structure/{upload.json()['id']}", headers=headers)

    assert response.status_code == 422
    assert response.json()["detail"] == "LLM response was not valid resume structure JSON"
    resume_file = db_session.get(ResumeFile, upload.json()["id"])
    assert resume_file.status == "structure_failed"
    assert resume_file.structure_error == "LLM response was not valid resume structure JSON"
    assert db_session.query(Experience).count() == 0
    assert db_session.query(Skill).count() == 0


def test_resume_structure_preserves_existing_sections_when_llm_omits_them(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)
    client.patch(
        "/api/profile",
        headers=headers,
        json={
            "projects": [
                {
                    "name": "Course Scheduler",
                    "description": "Built before resume import",
                    "tech_stack": ["React"],
                    "achievements": ["Used by 30 students"],
                }
            ],
            "skills": [{"name": "React", "category": "Frontend", "proficiency": "advanced"}],
        },
    )
    llm_client = StubLLMClient(
        [
            "{",
            '"profile":{"headline":"Backend engineer"},',
            '"experience":[{"company":"Acme","role":"Backend Intern","time":"Summer 2025",',
            '"description":"Built APIs","achievements":["Cut latency by 30%"]}],',
            '"follow_up_questions":["What product impact should we emphasize?"]',
            "}",
        ]
    )
    monkeypatch.setattr(resume_api, "build_llm_client", lambda settings: llm_client)
    upload = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", make_pdf_bytes("Alex Chen Backend Intern Acme"), "application/pdf")},
    )

    response = client.post(f"/api/resume/structure/{upload.json()['id']}", headers=headers)

    assert response.status_code == 200
    profile = db_session.query(Profile).one()
    assert profile.headline == "Backend engineer"
    assert db_session.query(Experience).one().company == "Acme"
    assert db_session.query(Project).one().name == "Course Scheduler"
    assert {skill.name for skill in db_session.query(Skill).all()} == {"React"}


def test_resume_structure_endpoint_rejects_files_owned_by_another_user(client, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    owner_headers = auth_headers(client)
    upload = client.post(
        "/api/resume/upload",
        headers=owner_headers,
        files={"file": ("resume.pdf", make_pdf_bytes("Private resume"), "application/pdf")},
    )
    other = client.post(
        "/api/auth/register",
        json={"email": "jamie.structure@example.com", "username": "jamiestructure", "password": "secret123"},
    ).json()

    response = client.post(
        f"/api/resume/structure/{upload.json()['id']}",
        headers={"Authorization": f"Bearer {other['access_token']}"},
    )

    assert response.status_code == 404


@pytest.mark.parametrize("endpoint", ["parse-status", "parsed"])
def test_resume_parse_endpoints_reject_files_owned_by_another_user(client, tmp_path, monkeypatch, endpoint):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    owner_headers = auth_headers(client)
    upload = client.post(
        "/api/resume/upload",
        headers=owner_headers,
        files={"file": ("resume.pdf", make_pdf_bytes("Private resume"), "application/pdf")},
    )
    other = client.post(
        "/api/auth/register",
        json={"email": "jamie.resume@example.com", "username": "jamieresume", "password": "secret123"},
    ).json()

    response = client.get(
        f"/api/resume/{endpoint}/{upload.json()['id']}",
        headers={"Authorization": f"Bearer {other['access_token']}"},
    )

    assert response.status_code == 404


@pytest.mark.parametrize("endpoint", ["parse-status", "parsed"])
def test_resume_parse_endpoints_require_authentication(client, tmp_path, monkeypatch, endpoint):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    owner_headers = auth_headers(client)
    upload = client.post(
        "/api/resume/upload",
        headers=owner_headers,
        files={"file": ("resume.pdf", make_pdf_bytes("Private resume"), "application/pdf")},
    )

    response = client.get(f"/api/resume/{endpoint}/{upload.json()['id']}")

    assert response.status_code == 401


def test_resume_upload_rejects_spoofed_pdf_content(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)

    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.pdf", b"not really a pdf", "application/pdf")},
    )

    assert response.status_code == 415
    assert response.json()["detail"] == "Unsupported resume file type. Upload a valid PDF or DOCX file."
    assert not (tmp_path / "resumes").exists()
    assert db_session.query(ResumeFile).count() == 0


def test_resume_upload_rejects_spoofed_docx_content(client, db_session, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    headers = auth_headers(client)
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.docx", b"PK\x03\x04not-docx", content_type)},
    )

    assert response.status_code == 415
    assert response.json()["detail"] == "Unsupported resume file type. Upload a valid PDF or DOCX file."
    assert not (tmp_path / "resumes").exists()
    assert db_session.query(ResumeFile).count() == 0


def test_resume_storage_stops_reading_after_size_limit(tmp_path):
    class ChunkedFile:
        def __init__(self):
            self.read_calls = 0

        async def read(self, size):
            self.read_calls += 1
            if self.read_calls == 1:
                return b"%PDF-1.4\n"
            if self.read_calls == 2:
                return b"x" * 8
            raise AssertionError("reader continued after upload exceeded the byte limit")

    upload = ChunkedFile()

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(store_validated_resume(upload, tmp_path / "resume.pdf", max_upload_bytes=10, suffix=".pdf"))

    assert exc_info.value.status_code == 413
    assert not (tmp_path / "resume.pdf").exists()


def test_resume_upload_removes_stored_file_when_db_commit_fails(client, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    stored_paths = []

    class FailingDb:
        rolled_back = False

        def add(self, resume_file):
            stored_paths.append(Path(resume_file.storage_path))

        def flush(self):
            pass

        def refresh(self, resume_file):
            pass

        def commit(self):
            raise RuntimeError("commit failed")

        def rollback(self):
            self.rolled_back = True

        def close(self):
            pass

    failing_db = FailingDb()

    def override_db():
        yield failing_db

    client.app.dependency_overrides[resume_api.get_current_user] = lambda: SimpleNamespace(id="user-1")
    client.app.dependency_overrides[resume_api.get_db] = override_db
    try:
        with pytest.raises(RuntimeError, match="commit failed"):
            client.post(
                "/api/resume/upload",
                files={"file": ("resume.pdf", b"%PDF-1.4\nsample resume", "application/pdf")},
            )
    finally:
        client.app.dependency_overrides.pop(resume_api.get_current_user, None)
        client.app.dependency_overrides.pop(resume_api.get_db, None)

    assert failing_db.rolled_back is True
    assert stored_paths
    assert not stored_paths[0].exists()


def test_resume_upload_removes_stored_file_when_db_refresh_fails_before_commit(client, tmp_path, monkeypatch):
    configure_resume_storage(monkeypatch, tmp_path / "resumes")
    stored_paths = []

    class FailingDb:
        committed = False
        rolled_back = False

        def add(self, resume_file):
            stored_paths.append(Path(resume_file.storage_path))

        def flush(self):
            pass

        def refresh(self, resume_file):
            raise RuntimeError("refresh failed")

        def commit(self):
            self.committed = True

        def rollback(self):
            self.rolled_back = True

        def close(self):
            pass

    failing_db = FailingDb()

    def override_db():
        yield failing_db

    client.app.dependency_overrides[resume_api.get_current_user] = lambda: SimpleNamespace(id="user-1")
    client.app.dependency_overrides[resume_api.get_db] = override_db
    try:
        with pytest.raises(RuntimeError, match="refresh failed"):
            client.post(
                "/api/resume/upload",
                files={"file": ("resume.pdf", b"%PDF-1.4\nsample resume", "application/pdf")},
            )
    finally:
        client.app.dependency_overrides.pop(resume_api.get_current_user, None)
        client.app.dependency_overrides.pop(resume_api.get_db, None)

    assert failing_db.committed is False
    assert failing_db.rolled_back is True
    assert stored_paths
    assert not stored_paths[0].exists()
